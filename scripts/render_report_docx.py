#!/usr/bin/env python3
from pathlib import Path
import argparse, json, re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SEVERITY_ORDER={"critical":0,"high":1,"medium":2,"low":3,"informational":4}

def clean_name(value):
    value=(value or "it-stod").strip().lower()
    value=re.sub(r"[^a-z0-9åäö]+","-",value,flags=re.I).strip("-")
    return value or "it-stod"

def set_cell_shading(cell, fill='D9E1F2'):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def prevent_row_split(row):
    trPr=row._tr.get_or_add_trPr(); cant=OxmlElement('w:cantSplit'); trPr.append(cant)

def set_cell_width(cell, width_cm):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.first_child_found_in('w:tcW')
    if tcW is None:
        tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm*567))); tcW.set(qn('w:type'),'dxa')

def set_cell_text(cell, text, bold=False):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(9)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    table=doc.add_table(rows=1, cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'; table.autofit=False if widths else True
    hdr=table.rows[0]; set_repeat_table_header(hdr); prevent_row_split(hdr)
    for i,h in enumerate(headers):
        set_cell_text(hdr.cells[i],h,True); set_cell_shading(hdr.cells[i])
        if widths: set_cell_width(hdr.cells[i], widths[i])
    for row in rows:
        r=table.add_row(); prevent_row_split(r); cells=r.cells
        for i,v in enumerate(row):
            set_cell_text(cells[i], v if v not in (None,'') else '–')
            if widths: set_cell_width(cells[i], widths[i])
    return table

def add_bullets(doc, items, empty='Inga identifierade poster.'):
    if not items:
        doc.add_paragraph(empty); return
    for item in items: doc.add_paragraph(str(item), style='List Bullet')

def add_heading(doc, text, level):
    p=doc.add_heading(text, level=level)
    return p

def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=paragraph.add_run('Sida ')
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text='PAGE'
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def setup_document(doc, title):
    sec=doc.sections[0]; sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)
    styles=doc.styles
    styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9.5)
    for s,size in [('Title',24),('Heading 1',16),('Heading 2',12.5),('Heading 3',10.5)]:
        styles[s].font.name='Aptos'; styles[s].font.size=Pt(size)
    header=sec.header.paragraphs[0]; header.text=title; header.runs[0].font.size=Pt(8)
    add_page_number(sec.footer.paragraphs[0])

def render(report):
    doc=Document(); m=report['metadata']; e=report['executive_summary']; s=report['scope']; sysov=report['system_overview']; arch=report['architecture_security']; cov=report['coverage']; rr=report['residual_risk']; app=report.get('appendix') or {}
    setup_document(doc,m['title'])
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(m['title'])
    if m.get('system_name'):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(m['system_name']); r.bold=True; r.font.size=Pt(14)
    doc.add_paragraph('')

    add_heading(doc,'Metadata',1)
    add_table(doc,['Fält','Värde'],[
        ['System/IT-stöd',m.get('system_name') or 'Ej angivet'],['Granskningsdatum',m.get('review_date') or 'Ej angivet'],['Granskningsläge',m.get('review_mode')],['Version',m.get('version')],['Underlagsreferens',m.get('source_reference') or 'Ej angivet']])

    add_heading(doc,'Sammanfattning',1); doc.add_paragraph(e['overall_assessment'])
    for label,key in [('Viktigaste fynd','key_findings'),('Viktigaste osäkerheter','key_uncertainties'),('Rekommenderade nästa steg','next_steps')]: add_heading(doc,label,2); add_bullets(doc,e[key])

    add_heading(doc,'Scope och analyserat underlag',1); doc.add_paragraph(s['requested_scope']); add_heading(doc,'Analyserat underlag',2); add_bullets(doc,s['reviewed_material']); add_heading(doc,'Avgränsningar',2); add_bullets(doc,s['limitations'])

    add_heading(doc,'System- och tekniköversikt',1)
    for key,label in [('frontend','Frontend'),('backend','Backend'),('data_stores','Datalager'),('integrations','Integrationer'),('deployment','Deployment')]: add_heading(doc,label,2); add_bullets(doc,sysov.get(key,[]))

    add_heading(doc,'Arkitekturell säkerhetsbild',1)
    for key,label in [('trust_boundaries','Trust boundaries'),('authentication_points','Autentiseringspunkter'),('authorization_points','Auktoriseringspunkter'),('administrative_interfaces','Administrativa gränssnitt'),('sensitive_data_flows','Känsliga dataflöden'),('observations','Observationer')]: add_heading(doc,label,2); add_bullets(doc,arch.get(key,[]))

    findings=sorted(report.get('findings',[]),key=lambda f:(SEVERITY_ORDER.get(f['severity'],99),f['id']))
    add_heading(doc,'Fynd',1)
    if findings:
        add_table(doc,['ID','Severity','Titel'],[[f['id'],f['severity'],f['title']] for f in findings], widths=[1.8,2.3,12.8])
        for f in findings:
            add_heading(doc,f"{f['id']} – {f['title']}",2)
            p=doc.add_paragraph()
            fields=[('Kategori',f['category']),('Severity',f['severity']),('Confidence',f['confidence']),('Status',f['status']),('Komponent',f.get('component') or 'Ej angivet'),('Manuell verifiering',f['manual_verification'])]
            for i,(label,value) in enumerate(fields):
                r=p.add_run(f"{label}: " ); r.bold=True; p.add_run(str(value))
                if i < len(fields)-1: p.add_run('  |  ')
            add_heading(doc,'Observation',3); doc.add_paragraph(f['observation'])
            if f.get('impact'): add_heading(doc,'Möjlig konsekvens',3); doc.add_paragraph(f['impact'])
            if f.get('reasoning'): add_heading(doc,'Resonemang',3); doc.add_paragraph(f['reasoning'])
            add_heading(doc,'Rekommenderad åtgärd',3); doc.add_paragraph(f['recommendation'])
            ev=f.get('evidence_details') or []
            if ev:
                add_heading(doc,'Evidens',3)
                for item in ev:
                    loc=f" ({item.get('location')})" if item.get('location') else ''
                    p=doc.add_paragraph(style='List Bullet'); r=p.add_run(f"{item['source']}{loc}: "); r.bold=True; p.add_run(item['description'])
            elif f.get('evidence'): add_heading(doc,'Evidens',3); add_bullets(doc,f['evidence'])
            if f.get('references'): add_heading(doc,'Referenser',3); add_bullets(doc,f['references'])
    else: doc.add_paragraph('Inga identifierade fynd.')

    add_heading(doc,'Coverage',1)
    for label,key in [('Granskat','reviewed'),('Ej granskat','not_reviewed'),('Ej verifierbart','not_verifiable')]: add_heading(doc,label,2); add_bullets(doc,cov[key])

    add_heading(doc,'Rekommenderade åtgärder',1)
    actions=report.get('recommended_actions',[])
    if actions: add_table(doc,['Prioritet','Åtgärd','Motivering','Relaterade fynd'],[[a['priority'],a['action'],a['reason'],', '.join(a.get('related_findings',[])) or '–'] for a in actions], widths=[2.1,6.4,5.8,2.6])
    else: doc.add_paragraph('Inga ytterligare åtgärder identifierade.')

    add_heading(doc,'Rekommenderad fortsatt granskning',1)
    follow=report.get('follow_up_review',[])
    if follow: add_table(doc,['Typ','Prioritet','Scope','Motivering','Verifieringsmål'],[[a['type'],a['priority'],a['scope'],a['reason'],a['verification_goal']] for a in follow], widths=[3.0,2.2,3.5,4.2,4.2])
    else: doc.add_paragraph('Ingen ytterligare särskild granskning rekommenderas.')

    add_heading(doc,'Kvarvarande risk',1); doc.add_paragraph(rr['summary'])
    for label,key in [('Från identifierade fynd','from_findings'),('Från ej verifierbara områden','from_unverified'),('Från områden utanför scope','from_out_of_scope')]: add_heading(doc,label,2); add_bullets(doc,rr[key])

    add_heading(doc,'Bilaga',1)
    for label,key in [('Metod','method'),('Evidensregister','evidence_register'),('Använda profiler','profiles_used'),('Referenser','references')]: add_heading(doc,label,2); add_bullets(doc,app.get(key,[]))
    return doc

def main():
    ap=argparse.ArgumentParser(); ap.add_argument('input_json'); ap.add_argument('-o','--output'); args=ap.parse_args()
    report=json.loads(Path(args.input_json).read_text(encoding='utf-8'))
    output=Path(args.output) if args.output else Path(f"sakerhetsgranskning-{clean_name(report['metadata'].get('system_name'))}.docx")
    render(report).save(output); print(output)
if __name__=='__main__': main()
